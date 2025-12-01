"""
元数据提取服务

从各种文件类型中提取基础和扩展元数据信息。
"""

import os
import hashlib
import mimetypes
from pathlib import Path
from typing import Dict, Any, Optional, List, Union
from datetime import datetime
import logging

# 文件处理库导入
try:
    from PyPDF2 import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logging.warning("PyPDF2未安装，PDF元数据提取功能不可用")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx未安装，Word文档元数据提取功能不可用")

try:
    from openpyxl import load_workbook
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    logging.warning("openpyxl未安装，Excel文档元数据提取功能不可用")

try:
    from pptx import Presentation
    PPTX_AVAILABLE = True
except ImportError:
    PPTX_AVAILABLE = False
    logging.warning("python-pptx未安装，PowerPoint文档元数据提取功能不可用")

try:
    from mutagen import File as MutagenFile
    AUDIO_AVAILABLE = True
except ImportError:
    AUDIO_AVAILABLE = False
    logging.warning("mutagen未安装，音频元数据提取功能不可用")

try:
    from PIL import Image
    from PIL.ExifTags import TAGS
    IMAGE_AVAILABLE = True
except ImportError:
    IMAGE_AVAILABLE = False
    logging.warning("Pillow未安装，图像元数据提取功能不可用")

logger = logging.getLogger(__name__)


class MetadataExtractor:
    """元数据提取器

    支持从各种文件类型中提取：
    - 基础元数据：文件名、大小、时间等
    - 扩展元数据：文档属性、音视频元数据、图像EXIF等
    - 文件哈希值：用于内容去重和变更检测
    """

    def __init__(self):
        """初始化元数据提取器"""
        self.supported_formats = {
            'document': ['.pdf', '.doc', '.docx', '.txt', '.md', '.rtf'],
            'spreadsheet': ['.xls', '.xlsx', '.csv'],
            'presentation': ['.ppt', '.pptx'],
            'image': ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.svg'],
            'audio': ['.mp3', '.wav', '.flac', '.aac', '.ogg', '.m4a'],
            'video': ['.mp4', '.avi', '.mkv', '.mov', '.wmv', '.flv', '.webm'],
            'archive': ['.zip', '.rar', '.7z', '.tar', '.gz']
        }

    def extract_metadata(self, file_input: Union[str, Any]) -> Dict[str, Any]:
        """提取文件元数据

        Args:
            file_input: 文件路径字符串或FileInfo对象

        Returns:
            Dict[str, Any]: 包含文件元数据的字典
        """
        try:
            # 处理不同的输入类型
            if hasattr(file_input, 'path'):  # FileInfo对象
                file_path = file_input.path
            else:  # 字符串路径
                file_path = file_input

            path = Path(file_path)
            if not path.exists():
                raise FileNotFoundError(f"文件不存在: {file_path}")

            # 基础元数据
            metadata = self._extract_basic_metadata(path)

            # 根据文件类型提取扩展元数据
            file_type = self._get_file_type(path.suffix.lower())
            metadata['file_type'] = file_type

            if file_type == 'document':
                metadata.update(self._extract_document_metadata(path))
            elif file_type == 'spreadsheet':
                metadata.update(self._extract_spreadsheet_metadata(path))
            elif file_type == 'presentation':
                metadata.update(self._extract_presentation_metadata(path))
            elif file_type == 'image':
                metadata.update(self._extract_image_metadata(path))
            elif file_type == 'audio':
                metadata.update(self._extract_audio_metadata(path))
            elif file_type == 'video':
                metadata.update(self._extract_video_metadata(path))

            # 计算文件哈希
            metadata['content_hash'] = self._calculate_file_hash(file_path)

            return metadata

        except Exception as e:
            logger.error(f"提取文件元数据失败 {file_path}: {e}")
            return {'error': str(e)}

    def _extract_basic_metadata(self, path: Path) -> Dict[str, Any]:
        """提取基础文件元数据"""
        try:
            stat = path.stat()
            mime_type, encoding = mimetypes.guess_type(str(path))

            return {
                'file_name': path.name,
                'file_path': str(path.absolute()),
                'file_extension': path.suffix.lower(),
                'file_size': stat.st_size,
                'file_size_mb': round(stat.st_size / 1024 / 1024, 2),
                'created_time': datetime.fromtimestamp(stat.st_ctime),
                'modified_time': datetime.fromtimestamp(stat.st_mtime),
                'accessed_time': datetime.fromtimestamp(stat.st_atime),
                'mime_type': mime_type or 'application/octet-stream',
                'encoding': encoding
            }
        except Exception as e:
            logger.error(f"提取基础元数据失败 {path}: {e}")
            return {}

    def _extract_document_metadata(self, path: Path) -> Dict[str, Any]:
        """提取文档元数据"""
        extension = path.suffix.lower()
        metadata = {}

        try:
            if extension == '.pdf' and PDF_AVAILABLE:
                metadata.update(self._extract_pdf_metadata(path))
            elif extension in ['.docx'] and DOCX_AVAILABLE:
                metadata.update(self._extract_docx_metadata(path))
            elif extension in ['.txt', '.md']:
                metadata.update(self._extract_text_metadata(path))

        except Exception as e:
            logger.error(f"提取文档元数据失败 {path}: {e}")
            metadata['extract_error'] = str(e)

        return metadata

    def _extract_pdf_metadata(self, path: Path) -> Dict[str, Any]:
        """提取PDF文档元数据"""
        metadata = {}

        try:
            with open(path, 'rb') as file:
                reader = PdfReader(file)

                # PDF信息
                if reader.metadata:
                    metadata.update({
                        'title': reader.metadata.get('/Title', ''),
                        'author': reader.metadata.get('/Author', ''),
                        'subject': reader.metadata.get('/Subject', ''),
                        'creator': reader.metadata.get('/Creator', ''),
                        'producer': reader.metadata.get('/Producer', ''),
                        'creation_date': reader.metadata.get('/CreationDate', ''),
                        'modification_date': reader.metadata.get('/ModDate', '')
                    })

                # PDF结构信息
                pdf_header = getattr(reader, 'pdf_header', None)
                pdf_version = 'unknown'
                if isinstance(pdf_header, dict):
                    pdf_version = pdf_header.get('version', 'unknown')
                elif isinstance(pdf_header, str):
                    # 如果pdf_header是字符串，尝试解析版本信息
                    if '%' in pdf_header:
                        try:
                            pdf_version = pdf_header.split('%')[1].strip().split('-')[0]
                        except:
                            pdf_version = pdf_header.strip()
                    else:
                        pdf_version = pdf_header.strip()

                metadata.update({
                    'page_count': len(reader.pages),
                    'is_encrypted': reader.is_encrypted,
                    'pdf_version': pdf_version
                })

        except Exception as e:
            logger.error(f"提取PDF元数据失败 {path}: {e}")

        return metadata

    def _extract_docx_metadata(self, path: Path) -> Dict[str, Any]:
        """提取Word文档元数据"""
        metadata = {}

        try:
            doc = Document(str(path))
            core_props = doc.core_properties

            metadata.update({
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'keywords': core_props.keywords or '',
                'category': core_props.category or '',
                'comments': core_props.comments or '',
                'created': core_props.created,
                'modified': core_props.modified,
                'last_modified_by': core_props.last_modified_by or '',
                'revision': core_props.revision or 0,
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables)
            })

        except Exception as e:
            logger.error(f"提取Word文档元数据失败 {path}: {e}")

        return metadata

    def _extract_text_metadata(self, path: Path) -> Dict[str, Any]:
        """提取文本文件元数据"""
        metadata = {}

        try:
            # 尝试不同编码读取文件
            encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
            content = None
            encoding_used = None

            for encoding in encodings:
                try:
                    with open(path, 'r', encoding=encoding) as file:
                        content = file.read()
                        encoding_used = encoding
                        break
                except UnicodeDecodeError:
                    continue

            if content:
                lines = content.split('\n')
                metadata.update({
                    'encoding': encoding_used,
                    'line_count': len(lines),
                    'character_count': len(content),
                    'word_count': len(content.split()),
                    'non_empty_line_count': len([line for line in lines if line.strip()])
                })

        except Exception as e:
            logger.error(f"提取文本文件元数据失败 {path}: {e}")

        return metadata

    def _extract_spreadsheet_metadata(self, path: Path) -> Dict[str, Any]:
        """提取Excel文档元数据"""
        metadata = {}

        try:
            if path.suffix.lower() in ['.xlsx'] and EXCEL_AVAILABLE:
                workbook = load_workbook(str(path), read_only=True)

                metadata.update({
                    'sheet_count': len(workbook.sheetnames),
                    'sheet_names': workbook.sheetnames,
                    'active_sheet': workbook.active.title if workbook.active else None
                })

                # 统计每个工作表的数据
                sheet_data = {}
                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    sheet_data[sheet_name] = {
                        'max_row': sheet.max_row,
                        'max_column': sheet.max_column
                    }

                metadata['sheet_data'] = sheet_data

        except Exception as e:
            logger.error(f"提取Excel文档元数据失败 {path}: {e}")

        return metadata

    def _extract_presentation_metadata(self, path: Path) -> Dict[str, Any]:
        """提取PowerPoint文档元数据"""
        metadata = {}

        try:
            if path.suffix.lower() in ['.pptx'] and PPTX_AVAILABLE:
                prs = Presentation(str(path))

                # 核心属性
                core_props = prs.core_properties
                metadata.update({
                    'title': core_props.title or '',
                    'author': core_props.author or '',
                    'subject': core_props.subject or '',
                    'created': core_props.created,
                    'modified': core_props.modified,
                    'slide_count': len(prs.slides)
                })

        except Exception as e:
            logger.error(f"提取PowerPoint文档元数据失败 {path}: {e}")

        return metadata

    def _extract_image_metadata(self, path: Path) -> Dict[str, Any]:
        """提取图像元数据"""
        metadata = {}

        try:
            if IMAGE_AVAILABLE:
                with Image.open(path) as img:
                    metadata.update({
                        'format': img.format,
                        'mode': img.mode,
                        'width': img.width,
                        'height': img.height,
                        'has_transparency': img.mode in ('RGBA', 'LA') or 'transparency' in img.info
                    })

                    # EXIF数据
                    if hasattr(img, '_getexif') and img._getexif():
                        exif_data = {}
                        exif = img._getexif()
                        for tag_id, value in exif.items():
                            tag = TAGS.get(tag_id, tag_id)
                            exif_data[tag] = str(value)
                        metadata['exif'] = exif_data

                    # 图像信息
                    metadata['info'] = dict(img.info)

        except Exception as e:
            logger.error(f"提取图像元数据失败 {path}: {e}")

        return metadata

    def _extract_audio_metadata(self, path: Path) -> Dict[str, Any]:
        """提取音频元数据"""
        metadata = {}

        try:
            if AUDIO_AVAILABLE:
                audio_file = MutagenFile(str(path))
                if audio_file:
                    metadata.update({
                        'duration': getattr(audio_file.info, 'length', 0),
                        'bitrate': getattr(audio_file.info, 'bitrate', 0),
                        'sample_rate': getattr(audio_file.info, 'sample_rate', 0),
                        'channels': getattr(audio_file.info, 'channels', 0),
                        'format': audio_file.mime[0] if audio_file.mime else 'unknown'
                    })

                    # 标签信息
                    tags = {}
                    if hasattr(audio_file, 'tags') and audio_file.tags:
                        for key, value in audio_file.tags.items():
                            if isinstance(value, list) and value:
                                tags[key] = str(value[0])
                            else:
                                tags[key] = str(value)

                    metadata['tags'] = tags

        except Exception as e:
            logger.error(f"提取音频元数据失败 {path}: {e}")

        return metadata

    def _extract_video_metadata(self, path: Path) -> Dict[str, Any]:
        """提取视频元数据"""
        # 注意：视频元数据提取需要额外的库如opencv-python或ffmpeg
        metadata = {}

        try:
            # 基础视频信息（仅通过文件大小和扩展名推测）
            metadata.update({
                'format': 'video',
                'note': '详细视频元数据提取需要安装opencv-python'
            })

        except Exception as e:
            logger.error(f"提取视频元数据失败 {path}: {e}")

        return metadata

    def _get_file_type(self, extension: str) -> str:
        """根据扩展名获取文件类型"""
        for file_type, extensions in self.supported_formats.items():
            if extension in extensions:
                return file_type
        return 'unknown'

    def _calculate_file_hash(self, file_path: str) -> str:
        """计算文件哈希值"""
        hash_md5 = hashlib.md5()

        try:
            with open(file_path, "rb") as f:
                # 分块读取，避免大文件内存问题
                for chunk in iter(lambda: f.read(4096), b""):
                    hash_md5.update(chunk)
            return hash_md5.hexdigest()
        except Exception as e:
            logger.error(f"计算文件哈希失败 {file_path}: {e}")
            return ""

    def get_supported_formats(self) -> Dict[str, List[str]]:
        """获取支持的文件格式"""
        return self.supported_formats.copy()
